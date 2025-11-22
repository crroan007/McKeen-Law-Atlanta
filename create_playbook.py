from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(51, 51, 51)

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('McKeen Law')
title_run.font.size = Pt(48)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_before = Pt(6)
subtitle_run = subtitle.add_run('SEO & AI Optimization Playbook')
subtitle_run.font.size = Pt(32)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

# Spacer
for _ in range(4):
    doc.add_paragraph()

# Tagline
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_run = tagline.add_run('Comprehensive Strategy for Organic Growth and AI-Driven Visibility')
tagline_run.font.size = Pt(14)
tagline_run.font.italic = True
tagline_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# Date/footer
footer_text = doc.add_paragraph()
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_text.add_run('Atlanta, Georgia | 2025')
footer_run.font.size = Pt(11)
footer_run.font.color.rgb = RGBColor(100, 100, 100)

add_page_break(doc)

# Executive Summary
p = doc.add_heading('Executive Summary', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

summary_text = """This playbook outlines McKeen Law's comprehensive SEO and AI optimization strategy designed to drive qualified lead generation through organic search, improve search visibility across core practice areas, and optimize presence across AI-powered search interfaces.

The strategy is organized into ten sequential phases spanning setup, technical optimization, information architecture, content development, and ongoing refinement. By executing these phases systematically, McKeen Law will establish market authority in Atlanta-area DUI defense while capturing high-intent keywords across county and municipal jurisdictions.

Key focus areas include Core Web Vitals optimization, conversion-optimized page design, trust-building through social proof and structured data, and active Google Business Profile management. The strategy prioritizes needle-moving initiatives that deliver measurable impact on call volume, form submissions, and chat engagement."""

p = doc.add_paragraph(summary_text)
p.paragraph_format.line_spacing = 1.15

doc.add_paragraph()

# Performance Objectives
p = doc.add_heading('Performance Objectives', level=2)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

objectives = [
    ('Conversion Volume', 'Increase qualified calls, form submissions, and chat interactions from organic channels'),
    ('Search Visibility', 'Achieve dominant market share on "Atlanta DUI lawyer/attorney" and county-specific DUI defense queries'),
    ('Click-Through Rate', 'Improve SERP CTR through optimized titles, meta descriptions, and rich result eligibility'),
    ('Trust & Authority', 'Build E-E-A-T signals through content quality, schema markup, review management, and attorney credentials'),
    ('Technical Excellence', 'Maintain Core Web Vitals below thresholds and ensure mobile-first user experience'),
]

for obj_title, obj_desc in objectives:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(obj_title + ': ')
    p_run.bold = True
    p.add_run(obj_desc)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Key Performance Indicators
p = doc.add_heading('Key Performance Indicators', level=2)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

kpi_list = [
    'GA4 Conversions: Calls, form submissions, and chat initiations from organic traffic',
    'Google Search Console: Click volume and CTR by query and query category',
    'Core Web Vitals: LCP, CLS, and INP metrics from real-user data (CrUX)',
    'Google Business Profile: Review volume, action requests, and Q&A engagement',
    'Organic Rankings: Position tracking for 50+ priority keywords across service lines and locations',
    'Organic Traffic: Trending traffic volume and user session quality metrics',
]

for kpi in kpi_list:
    p = doc.add_paragraph(kpi, style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)

add_page_break(doc)

# Phase 0
p = doc.add_heading('Phase 0: Foundation Setup', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 1–2')

p = doc.add_heading('Objectives', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

phase0_items = [
    'Establish proper site configuration and indexation controls',
    'Implement measurement infrastructure for tracking and analytics',
    'Enforce URL consistency and canonical structure',
]

for item in phase0_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Key Initiatives', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

initiatives = [
    ('Google Search Console Verification', 'Verify domain ownership; monitor coverage and indexation; plan and submit XML sitemap with canonical and 200-status URLs only'),
    ('Google Analytics 4 Integration', 'Set up GA4 in a single, consolidated Google Tag Manager container; establish conversion goals for calls, forms, and chat'),
    ('URL Structure & Canonicalization', 'Enforce HTTPS; designate primary canonical host; establish consistent trailing slash convention; implement 301 redirects for old URLs'),
    ('Robots.txt Configuration', 'Allow full site crawl; disallow admin/private areas; point crawlers to XML sitemap'),
    ('URL Slug Strategy', 'Adopt clean, descriptive slugs (e.g., /dui-defense, /locations/fulton-county, /attorneys/john-smith)'),
]

for title, desc in initiatives:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

add_page_break(doc)

# Phase 1
p = doc.add_heading('Phase 1: Technical Baseline', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 2–3')

p = doc.add_heading('Objectives', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

phase1_objs = [
    'Establish technical SEO fundamentals',
    'Meet Core Web Vitals thresholds',
    'Ensure mobile-first responsive design',
    'Implement accessible markup',
]

for item in phase1_objs:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Core Web Vitals Targets', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

vitals = [
    ('Largest Contentful Paint (LCP)', '< 2.5 seconds', [
        'Preload hero images with correct width/height attributes',
        'Lazy-load images and media below the fold',
        'Minimize large third-party scripts that block rendering',
    ]),
    ('Cumulative Layout Shift (CLS)', '< 0.1', [
        'Set explicit width/height on media elements',
        'Reserve space for ads, embeds, and dynamic content',
        'Avoid inserting content above existing DOM elements',
    ]),
    ('Interaction to Next Paint (INP)', '< 200ms', [
        'Defer non-critical JavaScript execution',
        'Keep bundle sizes compact; code-split where appropriate',
        'Minimize third-party tag overhead',
    ]),
]

for vital, target, tactics in vitals:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(vital + ' ')
    run.bold = True
    p.add_run(target)
    for tactic in tactics:
        doc.add_paragraph(tactic, style='List Bullet 2')

p = doc.add_heading('Mobile & Accessibility', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

a11y_items = [
    'Responsive layout with no horizontal scrolling',
    'Touch targets (buttons, links) sized ≥ 48px for easy mobile interaction',
    'Single H1 per page; logical semantic heading hierarchy',
    'Form labels associated with input fields; ARIA labels for custom controls',
    'Tel: links for phone numbers to enable mobile click-to-call',
    'Visible focus states for keyboard navigation',
    'Alt text on all images; descriptive filenames for accessibility',
    'Sufficient color contrast (4.5:1 for body text)',
    'System or self-hosted fonts with font-display: swap to avoid FOIT',
]

for item in a11y_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Structured Data Foundation', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('Implement BreadcrumbList schema on all pages to enable breadcrumb navigation in search results', style='List Bullet')

add_page_break(doc)

# Phase 2
p = doc.add_heading('Phase 2: Information Architecture', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 3')

p = doc.add_heading('Primary Navigation Structure', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

nav_items = [
    'Home',
    'DUI Defense',
    'Charges & Defenses',
    'Case Results',
    'Reviews',
    'Locations',
    'Attorneys',
    'Contact',
]

p = doc.add_paragraph('Main navigation items (top bar):')
for item in nav_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Additional deep links positioned in footer and XML sitemap', style='List Bullet')

p = doc.add_heading('Location Architecture', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

location_arch = [
    ('Geographic Hub', '/locations/atlanta-dui-lawyer (primary metropolitan area)'),
    ('County Pages', 'Fulton, DeKalb, Cobb, Gwinnett, Cherokee, Henry (with unique local content)'),
    ('City Pages', 'Only when unique court information, local judges, or case results exist—avoid doorway page duplicates'),
]

for arch_type, desc in location_arch:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(arch_type + ': ')
    run.bold = True
    p.add_run(desc)

p = doc.add_heading('Content Hub Strategy', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

hubs = [
    'DUI Defense (core hub)',
    'License & Administrative License Suspension (ALS)',
    'Right to Refuse Breath/Blood Tests',
    'Field Sobriety Tests, Breath Tests, Blood Tests',
    'Implied Consent & Less-Safe Alternative',
    'Under-21 DUI Offenses',
    'CDL & Commercial DUI',
    'Boating Under the Influence (BUI)',
    'Drug-Related DUI',
    'Reckless Driving & Related Charges',
]

for hub in hubs:
    doc.add_paragraph(hub, style='List Bullet')

add_page_break(doc)

# Phase 3
p = doc.add_heading('Phase 3: Conversion & Trust Architecture', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 3–4')

p = doc.add_heading('Hero Section Design', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

hero_elements = [
    '24/7 click-to-call button with tel: link for mobile conversion',
    'Bold "Free Consultation" CTA button with high contrast',
    'Short contact form (name, phone, message) for non-call users',
    'Professional imagery or video reinforcing trust and competence',
]

for elem in hero_elements:
    doc.add_paragraph(elem, style='List Bullet')

p = doc.add_heading('Mobile Sticky Bar', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

sticky_desc = 'Fixed bottom bar on mobile devices with prominent call, text, and consultation buttons; high contrast colors; does not obscure critical content'
doc.add_paragraph(sticky_desc, style='List Bullet')

p = doc.add_heading('Trust & Authority Signals', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

trust_blocks = [
    'Case Results Statistics: Number of favorable outcomes, DUI dismissals, and reduced charges',
    'Review Snippets: Star ratings and client testimonial excerpts from verified reviews',
    'Professional Badges: State bar association membership, recognition, certifications',
    'Attorney Credentials: Bar license numbers, prosecutor background, years of experience',
    'Social Proof: Media features, speaking engagements, awards',
]

for block in trust_blocks:
    doc.add_paragraph(block, style='List Bullet')

p = doc.add_heading('Contact Page Essentials', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

contact_items = [
    'Name, Address, Phone (NAP) clearly displayed',
    'Async map embed with directions and parking information',
    'Hours of operation and holiday schedule',
    'Short contact form with clear labels',
    'Alternative contact methods (phone, text, email)',
]

for item in contact_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 4
p = doc.add_heading('Phase 4: Content Development', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Month 1–2')

p = doc.add_heading('Home Page', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

home_elements = [
    'Unique value proposition (USP) distinguishing McKeen Law from competitors',
    'CTA stack: multiple high-contrast calls-to-action for different user intents',
    'Proof blocks: case results, review aggregates, certifications',
    'Attorney snippet with photos, credentials, specialization',
    'Locations teaser with links to top jurisdictions',
    'FAQ snippet with schema for potential featured snippets',
]

for elem in home_elements:
    doc.add_paragraph(elem, style='List Bullet')

p = doc.add_heading('DUI Defense Hub', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

hub_content = [
    'Georgia DUI statutes and penalties overview',
    'DUI arrest and prosecution timeline',
    'Defense strategies and options',
    'Potential penalties and sentencing guidelines',
    'Client action checklist for DUI charges',
    'Comprehensive FAQ section',
]

for item in hub_content:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('County & City Pages', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

county_desc = 'Each page must contain unique, locally relevant content to avoid duplicate content penalties.'
doc.add_paragraph(county_desc)

doc.add_paragraph()

county_elements = [
    'Local introduction: jurisdiction-specific DUI defense overview',
    'Local courts and judges: relevant courthouse information',
    'Local DUI process: jurisdiction-specific procedures and timelines',
    'Local case result: specific outcome from that jurisdiction',
    'Local review: client testimonial specific to the location',
    'Async map embed with directions',
    'NAP and hours of operation',
    '2–3 internal links to hub pages and contact page',
]

for elem in county_elements:
    doc.add_paragraph(elem, style='List Bullet')

p = doc.add_heading('Attorney Profiles', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

attorney_items = [
    'Professional biography and background',
    'State bar license number and disciplinary record link',
    'Prosecutor background and litigation experience',
    'Media features and speaking engagements',
    'Linked case results and client reviews',
]

for item in attorney_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Case Results Page', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('Filterable case results by charge type, location, and year for user-driven discovery', style='List Bullet')

p = doc.add_heading('Reviews Page', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('Aggregate review ratings and source snippets from verified review platforms', style='List Bullet')

p = doc.add_heading('Blog & Content Cadence', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

blog_freq = 'Publish 2–4 articles per month on Georgia DUI updates, court procedure changes, and notable cases'
doc.add_paragraph(blog_freq)

doc.add_paragraph()

blog_strategy = [
    'Each post links to relevant hub and location pages',
    'Clear introduction with answer to main query',
    'Bullet points and short paragraphs for scanners',
    'Dated publication and regular update timestamps',
    'Cited sources for legal and statutory claims',
    'AI-friendly formatting: natural language answers, clear structure',
]

for item in blog_strategy:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 5
p = doc.add_heading('Phase 5: Structured Data & Schema Markup', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Month 1–2')

p = doc.add_heading('Schema Implementation Strategy', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

schema_types = [
    ('Organization', 'Firm name, logo, contact information, social media profiles'),
    ('LocalBusiness / LegalService', 'Name, Address, Phone (NAP), hours, services, aligned with GBP'),
    ('BreadcrumbList', 'Breadcrumb navigation on all pages'),
    ('FAQPage', 'FAQ sections for featured snippet eligibility'),
    ('Review / AggregateRating', 'Star ratings and reviews (policy-compliant and visible on page)'),
    ('Article / BlogPosting', 'Blog and news content with author, date, and images'),
    ('Speakable', 'Top FAQs and concise answers for voice search compatibility'),
]

for schema_type, desc in schema_types:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(schema_type + ': ')
    run.bold = True
    p.add_run(desc)

p = doc.add_heading('Validation & Maintenance', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

validation_items = [
    'Use Google Rich Results Test and Schema.org validation tools',
    'Remove schema markup for content not visibly displayed on the page',
    'Monitor Google Search Console for structured data errors',
    'Refresh and validate schema during content updates',
]

for item in validation_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 6
p = doc.add_heading('Phase 6: Google Business Profile Management', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 2–4; ongoing weekly management')

p = doc.add_heading('Profile Optimization', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

gbp_items = [
    ('Primary Category', 'DUI/Criminal Defense'),
    ('Secondary Categories', 'Relevant service categories (e.g., criminal law, legal services)'),
    ('Services List', 'Detailed services offered with prices/descriptions'),
    ('Hours & Holiday Hours', 'Accurate operating hours and holiday schedule'),
    ('Website & Appointment Links', 'UTM-tagged links for tracking GBP traffic'),
    ('Business Description', 'Unique value proposition and differentiators'),
]

for category, content in gbp_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(category + ': ')
    run.bold = True
    p.add_run(content)

p = doc.add_heading('Content & Engagement', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

engagement_items = [
    ('Photos', 'Post monthly photos of team, office, signage for social proof'),
    ('Posts', 'Weekly posts about case wins, practice tips, legal updates'),
    ('Q&A', 'Seed real FAQs; monitor and answer new questions promptly'),
    ('Reviews', 'Request reviews after wins and positive client interactions; respond naturally'),
]

for activity, desc in engagement_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(activity + ': ')
    run.bold = True
    p.add_run(desc)

p = doc.add_heading('NAP Consistency', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

nap_text = 'Ensure Name, Address, Phone consistency across website, schema markup, and GBP. Use async map embed on contact/location pages.'
doc.add_paragraph(nap_text, style='List Bullet')

add_page_break(doc)

# Phase 7
p = doc.add_heading('Phase 7: Analytics & Tracking Configuration', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Week 2–3')

p = doc.add_heading('Google Analytics 4 Setup', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

ga4_items = [
    'Consolidated Google Tag Manager container for all tracking',
    'Event tracking for key user interactions:',
]

for item in ga4_items:
    if item.endswith(':'):
        doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph(item, style='List Bullet')

events = [
    'Tel: link clicks (identify phone click intent)',
    'Form submissions (contact form, chat initiation)',
    'Sticky mobile bar interactions',
    'Chat widget activation',
    'Map clicks and directions requests',
    'Video playback',
    'Outbound clicks to review platforms',
]

for event in events:
    doc.add_paragraph(event, style='List Bullet 2')

doc.add_paragraph('Mark all events as conversions where applicable', style='List Bullet')

p = doc.add_heading('Call Tracking', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

tracking_items = [
    'Dynamic phone numbers on landing pages for call source attribution',
    'Single primary number in footer, schema, and GBP for brand consistency',
    'UTM parameters on all trackable links and GBP CTAs',
]

for item in tracking_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Data Validation', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

validation_items = [
    'Verify events in GA4 DebugView before live deployment',
    'Monitor GSC for crawl errors, coverage, and indexation status',
    'Track Core Web Vitals via CrUX and PageSpeed Insights',
]

for item in validation_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 8
p = doc.add_heading('Phase 8: Link Building & Authority', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Month 2–6; ongoing')

p = doc.add_heading('Internal Link Strategy', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

internal_strategy = [
    'Hub-to-subpage linking: connect DUI Defense hub to all relevant defense category pages',
    'Blog-to-hub linking: each blog post links to relevant hub and location pages',
    'Descriptive anchor text: use natural, keyword-relevant anchor text (avoid generic "click here")',
    'Orphan page elimination: ensure all pages receive at least 2–3 internal links from higher-authority pages',
]

for item in internal_strategy:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('External Link Acquisition', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

external_sources = [
    'State and local bar associations (links from official listings)',
    'Chamber of Commerce membership and directories',
    'Legal sponsorships and community involvement',
    'Guest columns in legal journals and publications',
    'Podcast interviews and legal commentary',
    'Local media features and news coverage',
]

for source in external_sources:
    doc.add_paragraph(source, style='List Bullet')

p = doc.add_heading('Link Integrity', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

integrity_items = [
    'Mark sponsored links with rel="sponsored"',
    'Mark user-generated content links with rel="ugc"',
    'Avoid link schemes, excessive footer links, or unnatural linking patterns',
]

for item in integrity_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 9
p = doc.add_heading('Phase 9: User Experience & Conversion Refinement', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Ongoing')

p = doc.add_heading('Performance Optimization', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

perf_items = [
    'Keep pages fast and lightweight: minimize HTTP requests, optimize images',
    'Reduce cumulative layout shift: reserve space for dynamic content',
    'Minimize interruptive popups and modals on first load',
]

for item in perf_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Conversion Testing', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

testing_items = [
    'A/B test CTA copy (e.g., "Call Now" vs. "Free Consultation")',
    'Test button colors and placement for maximum visibility',
    'Test form field count and progression for drop-off reduction',
    'Track conversion rates by template and user cohort',
]

for item in testing_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Design Standards', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

design_items = [
    'Maintain high-contrast, legible typography for readability',
    'Ensure forms are short, clear, and mobile-optimized',
    'Use consistent branding and visual hierarchy throughout',
]

for item in design_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Phase 10
p = doc.add_heading('Phase 10: Ongoing Maintenance & Monitoring', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p_run = p.add_run('Timeline: ')
p_run.bold = True
p.add_run('Monthly')

p = doc.add_heading('Content Refresh', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

refresh_items = [
    'Update top-performing pages with new case results and client reviews',
    'Prune thin city pages lacking unique content; consolidate with county pages where appropriate',
    'Refresh blog content with new legal developments and case updates',
]

for item in refresh_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Technical Audits', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

audit_items = [
    'Monitor 404 errors in GSC; implement appropriate 301 redirects',
    'Audit redirect chains; simplify to direct 301 redirects',
    'Validate structured data in GSC; fix or remove broken schema',
    'Trim unused tags and tracking code; maintain lean tag architecture',
    'Update XML sitemap with new and modified pages',
]

for item in audit_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Google Business Profile Maintenance', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

gbp_maintenance = [
    'Verify hours and holiday schedules are current',
    'Publish weekly posts with case wins and practice tips',
    'Monitor and answer Q&A submissions promptly',
    'Respond to all reviews—positive and negative—professionally',
]

for item in gbp_maintenance:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('Performance Tracking', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

tracking_items = [
    'Monitor keyword rankings for 50+ priority terms via GSC',
    'Track CTR trends by query category and seasonality',
    'Investigate ranking drops and traffic anomalies via GSC and GA4',
    'Monitor Core Web Vitals for regressions; fix performance issues promptly',
    'Review conversion trends (calls, forms, chat) and optimize underperforming pages',
]

for item in tracking_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# AI Surface Optimization
p = doc.add_heading('AI Surface Optimization Strategy', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

intro_text = 'As AI-powered search interfaces (Google Gemini, OpenAI ChatGPT, Grok) become prominent discovery channels, content optimization for these surfaces is critical. This section details how to optimize McKeen Law content for AI retrieval and citation.'
p = doc.add_paragraph(intro_text)
p.paragraph_format.line_spacing = 1.15

p = doc.add_heading('Content Structure for AI', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

ai_structure = [
    ('Concise Summaries', 'Lead each page with a 2–3 sentence summary answering the main question'),
    ('Clear Headings', 'Use semantic heading hierarchy (H1 → H2 → H3) for logical structure'),
    ('Bullet Points', 'Break content into scannable, digestible bullets for LLM token efficiency'),
    ('FAQ Blocks', 'Include FAQ sections with matching schema markup for featured snippet eligibility'),
]

for technique, desc in ai_structure:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(technique + ': ')
    run.bold = True
    p.add_run(desc)

p = doc.add_heading('Credibility Signals for AI', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

credibility_items = [
    ('Publication Date', 'Include clear "Published" and "Last Updated" dates'),
    ('Source Attribution', 'Cite specific Georgia DUI statutes, court cases, and legal references'),
    ('Author Information', 'Display attorney name, credentials, and bar license number'),
    ('Experience Signals', 'Mention years of practice, prosecutorial background, case results'),
]

for signal, desc in credibility_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(signal + ': ')
    run.bold = True
    p.add_run(desc)

p = doc.add_heading('Technical Requirements', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

tech_reqs = [
    'Server-rendered content: ensure primary content is in HTML, not hidden behind JavaScript',
    'No interstitials: avoid splash screens or ads that obscure content on first load',
    'Transcripts: provide text transcripts for video content',
    'Image descriptions: use descriptive alt text and filenames for media',
]

for req in tech_reqs:
    doc.add_paragraph(req, style='List Bullet')

p = doc.add_heading('Natural Language Optimization', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

nl_items = [
    'Answer common queries naturally: "What happens after a DUI arrest in Atlanta?" should have a clear, conversational answer',
    'Avoid keyword stuffing: write for human readers, not search algorithms',
    'Use natural question phrasing in FAQ sections',
]

for item in nl_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_heading('GBP Optimization for AI', level=3)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

gbp_ai = [
    'Keep Q&A current and accurate; AI relies heavily on GBP Q&A for factual information',
    'Ensure consistent Name, Address, Phone (NAP) across site, schema, and GBP',
    'Maintain accurate hours and holiday schedules',
]

for item in gbp_ai:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# Needle Movers
p = doc.add_heading('High-Impact Priorities: Biggest Needle Movers', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

intro = 'For teams with limited time or resources, prioritize these initiatives for maximum impact on organic lead generation:'
doc.add_paragraph(intro)

needles = [
    {
        'title': '1. Fast, Conversion-First Hero & Sticky CTAs',
        'items': [
            '24/7 click-to-call button with tel: link',
            'Mobile sticky bar with prominent call, text, consult buttons',
            'Short contact form for non-call users',
            'High-contrast colors for visibility',
        ]
    },
    {
        'title': '2. Unique, Locally Relevant County & City Pages',
        'items': [
            'No duplicate content; each page must have unique local information',
            'Local courts, judges, DUI procedures, and case results',
            'Proof of local expertise (case results, reviews from that jurisdiction)',
            'Links to hub and contact pages',
        ]
    },
    {
        'title': '3. Optimized Google Business Profile',
        'items': [
            'Accurate NAP, hours, services',
            'Weekly posts with case wins and practice tips',
            'Active Q&A monitoring and responses',
            'Review program with responses to all reviews',
        ]
    },
    {
        'title': '4. Strong DUI Defense Hub with FAQ Schema',
        'items': [
            'Comprehensive hub page covering statutes, process, defenses, penalties',
            'FAQ section with matching schema markup for featured snippets',
            'Internal links from subpages and blog posts',
        ]
    },
    {
        'title': '5. Core Web Vitals & Speed Discipline',
        'items': [
            'LCP < 2.5s: optimize images, preload hero assets',
            'CLS < 0.1: reserve space, avoid dynamic insertions',
            'INP < 200ms: defer non-critical JS, minimize third-party overhead',
        ]
    },
    {
        'title': '6. Lean, Unified Analytics',
        'items': [
            'Single GTM container and GA4 property',
            'Event tracking for key conversions: calls, forms, chat',
            'Call tracking with dynamic numbers on landing pages',
        ]
    },
]

for needle in needles:
    p = doc.add_paragraph(needle['title'], style='Heading 3')
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    for item in needle['items']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

add_page_break(doc)

# Appendix / Conclusion
p = doc.add_heading('Conclusion', level=1)
for run in p.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

conclusion = """This playbook provides McKeen Law with a comprehensive, phased strategy for organic growth. By systematically executing these ten phases—from foundation setup through ongoing maintenance—McKeen Law will establish market authority in Atlanta DUI defense, capture high-intent keywords, and drive qualified leads from organic search.

Success depends on consistent execution across technical SEO, content quality, trust-building, and conversion optimization. Regular monitoring of KPIs in GA4 and Google Search Console will enable quick identification and resolution of issues.

The biggest opportunity lies in conversion-optimized pages, locally relevant content, and active Google Business Profile management. These three elements deliver the highest ROI and should receive priority focus.

As search interfaces evolve—particularly the rise of AI-powered discovery—the content and structured data principles in this playbook will ensure McKeen Law remains visible and citable across all search channels."""

p = doc.add_paragraph(conclusion)
p.paragraph_format.line_spacing = 1.15

doc.save('SEO_Playbook_McKinsey_Style.docx')
print("Success: SEO_Playbook_McKinsey_Style.docx created")
